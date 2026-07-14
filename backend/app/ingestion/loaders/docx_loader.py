from pathlib import Path

import docx

from app.ingestion.loaders.base import RawSection


class DocxLoader:
    """Extracts paragraph-level text (and table cell text) from a .docx file."""

    def load(self, file_path: Path) -> list[RawSection]:
        document = docx.Document(str(file_path))
        sections: list[RawSection] = []

        for i, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            sections.append(
                RawSection(
                    text=text,
                    locator=f"paragraph {i}",
                    metadata={"style": paragraph.style.name if paragraph.style else None},
                )
            )

        for t, table in enumerate(document.tables):
            for r, row in enumerate(table.rows):
                row_text = " | ".join(cell.text.strip() for cell in row.cells).strip()
                if not row_text:
                    continue
                sections.append(
                    RawSection(
                        text=row_text,
                        locator=f"table {t} row {r}",
                        metadata={"style": "table"},
                    )
                )

        return sections
