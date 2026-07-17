import io
from collections.abc import Iterable

import docx
from sqlalchemy import select

from app.db.models import Document, DocumentRelationship
from app.db.session import async_session_factory
from app.ingestion.loaders.docx_loader import DocxLoader
from app.ingestion.relationships import (
    SectionAnnotationExtractor,
    extract_and_store_relationships,
    find_candidate_sections,
    reconcile_relationships,
)
from app.llm.interfaces import Message

PRECEDENCE_TEXT = (
    "This regional handbook is read together with the global Acme Employee "
    "Handbook. Where a conflict arises specifically with respect to PAID TIME "
    "OFF (PTO), the LOCAL PTO POLICY set out in this APAC Benefits Handbook "
    "TAKES PRECEDENCE over any conflicting PTO provision in the global handbook "
    "for employees covered by this document."
)

PRECEDENCE_JSON = """
{
  "relationships": [
    {
      "target_doc_ref": "global Acme Employee Handbook",
      "relation_type": "precedence",
      "topic": "PTO",
      "precedence": "source_over_target",
      "source_text": "the LOCAL PTO POLICY ... TAKES PRECEDENCE over any conflicting PTO provision"
    }
  ],
  "geographic_scope": null
}
"""


class StubLLMClient:
    def __init__(self, response: str) -> None:
        self._response = response

    def generate(self, messages: list[Message], stream: bool = True) -> Iterable[str]:
        yield self._response


def _make_apac_handbook_bytes() -> bytes:
    document = docx.Document()
    document.add_paragraph("APAC Benefits Handbook", style="Title")
    document.add_paragraph("Conflicts and Precedence", style="Heading 1")
    document.add_paragraph(PRECEDENCE_TEXT)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_find_candidate_sections_flags_precedence_section(tmp_path):
    docx_path = tmp_path / "apac.docx"
    docx_path.write_bytes(_make_apac_handbook_bytes())

    elements = DocxLoader().load(docx_path)
    candidates = find_candidate_sections(elements)

    assert len(candidates) == 1
    assert candidates[0].heading_path == ("Conflicts and Precedence",)


async def test_extract_and_store_relationships_creates_precedence_row(tmp_path):
    docx_path = tmp_path / "apac.docx"
    docx_path.write_bytes(_make_apac_handbook_bytes())
    elements = DocxLoader().load(docx_path)

    extractor = SectionAnnotationExtractor(StubLLMClient(PRECEDENCE_JSON))

    async with async_session_factory() as db:
        document = Document(
            filename="apac.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            storage_path=str(docx_path),
            title="APAC Benefits Handbook",
        )
        db.add(document)
        await db.flush()

        geo_overrides, hints = await extract_and_store_relationships(db, document, elements, extractor)
        await db.commit()

        assert geo_overrides == {}  # this section only stated a relationship, no geo scope
        assert len(hints) == 1
        assert hints[0].relation_type == "precedence"
        assert hints[0].topic == "PTO"
        assert hints[0].target_doc_ref == "global Acme Employee Handbook"

        result = await db.execute(
            select(DocumentRelationship).where(DocumentRelationship.source_doc_id == document.id)
        )
        relationships = result.scalars().all()
        assert len(relationships) == 1
        assert relationships[0].relation_type == "precedence"
        assert relationships[0].topic == "PTO"
        assert relationships[0].target_doc_id is None  # global handbook not uploaded yet


GENERAL_RULE_TEXT = (
    "Where the perks and benefits described in this handbook conflict with "
    "those described in another Acme policy or handbook, the MORE GENEROUS "
    "perk or benefit applies to you."
)

GENERAL_RULE_JSON = (
    '{"relationships": [], "geographic_scope": null, "default_precedence_rule": '
    '"The more generous perk or benefit applies when this handbook conflicts '
    'with another Acme policy or handbook."}'
)


def _make_global_handbook_bytes() -> bytes:
    document = docx.Document()
    document.add_paragraph("Acme Employee Handbook", style="Title")
    document.add_paragraph("Conflicts and Precedence", style="Heading 1")
    document.add_paragraph(GENERAL_RULE_TEXT)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def test_extract_and_store_relationships_sets_default_precedence_rule(tmp_path):
    docx_path = tmp_path / "global.docx"
    docx_path.write_bytes(_make_global_handbook_bytes())
    elements = DocxLoader().load(docx_path)

    extractor = SectionAnnotationExtractor(StubLLMClient(GENERAL_RULE_JSON))

    async with async_session_factory() as db:
        document = Document(
            filename="global.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            storage_path=str(docx_path),
            title="Acme Employee Handbook",
        )
        db.add(document)
        await db.flush()

        await extract_and_store_relationships(db, document, elements, extractor)

        assert document.default_precedence_rule == (
            "The more generous perk or benefit applies when this handbook "
            "conflicts with another Acme policy or handbook."
        )


async def test_resolve_target_prefers_is_latest_on_tie(tmp_path):
    # Regression test for a real bug found via manual/live testing: two
    # versions of "Acme Employee Handbook" (2025, 2026) score identically
    # against a reference like "global Acme Employee Handbook" -- comparing by
    # fuzzy score alone left the tie-break to whichever row the DB query
    # happened to return first, which in practice resolved to the superseded
    # 2025 document instead of the current 2026 one.
    docx_path = tmp_path / "apac.docx"
    docx_path.write_bytes(_make_apac_handbook_bytes())
    elements = DocxLoader().load(docx_path)

    extractor = SectionAnnotationExtractor(StubLLMClient(PRECEDENCE_JSON))

    async with async_session_factory() as db:
        older = Document(
            filename="handbook-2025.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            storage_path="/tmp/handbook-2025.docx",
            title="Acme Employee Handbook",
            is_latest=False,
        )
        newer = Document(
            filename="handbook-2026.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            storage_path="/tmp/handbook-2026.docx",
            title="Acme Employee Handbook",
            is_latest=True,
        )
        db.add_all([older, newer])
        await db.flush()

        source = Document(
            filename="apac.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            storage_path=str(docx_path),
            title="APAC Benefits Handbook",
        )
        db.add(source)
        await db.flush()

        await extract_and_store_relationships(db, source, elements, extractor)

        result = await db.execute(
            select(DocumentRelationship).where(DocumentRelationship.source_doc_id == source.id)
        )
        relationship = result.scalar_one()
        assert relationship.target_doc_id == newer.id


async def test_extract_and_store_relationships_leaves_default_precedence_rule_none_when_unstated(
    tmp_path,
):
    docx_path = tmp_path / "apac.docx"
    docx_path.write_bytes(_make_apac_handbook_bytes())
    elements = DocxLoader().load(docx_path)

    extractor = SectionAnnotationExtractor(StubLLMClient(PRECEDENCE_JSON))

    async with async_session_factory() as db:
        document = Document(
            filename="apac.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            storage_path=str(docx_path),
            title="APAC Benefits Handbook 2",
        )
        db.add(document)
        await db.flush()

        await extract_and_store_relationships(db, document, elements, extractor)

        assert document.default_precedence_rule is None


async def test_reconciliation_backfills_target_once_referenced_doc_uploaded():
    async with async_session_factory() as db:
        source = Document(
            filename="apac2.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            storage_path="/tmp/apac2.docx",
            title="APAC Benefits Handbook v2",
        )
        db.add(source)
        await db.flush()

        db.add(
            DocumentRelationship(
                source_doc_id=source.id,
                target_doc_id=None,
                target_doc_ref="global Acme Employee Handbook",
                relation_type="precedence",
                topic="PTO",
                source_text=PRECEDENCE_TEXT,
            )
        )
        await db.commit()

        # The referenced document arrives later.
        target = Document(
            filename="global.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            storage_path="/tmp/global.docx",
            title="Acme Employee Handbook",
        )
        db.add(target)
        await db.flush()

        await reconcile_relationships(db, target)
        await db.commit()

        result = await db.execute(
            select(DocumentRelationship).where(DocumentRelationship.source_doc_id == source.id)
        )
        relationship = result.scalar_one()
        assert relationship.target_doc_id == target.id
