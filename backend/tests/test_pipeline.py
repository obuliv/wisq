from collections.abc import Iterable

import docx
from sqlalchemy import select

from app.db.models import Document, DocumentRelationship
from app.db.session import async_session_factory
from app.ingestion.chunking import SectionAwareChunker
from app.ingestion.metadata import CompositeMetadataExtractor
from app.ingestion.pipeline import IngestionPipeline
from app.ingestion.relationships import SectionAnnotationExtractor
from app.llm.fakes import FakeLLMClient
from app.llm.interfaces import Message
from app.rag.fakes import FakeEmbedder, InMemoryVectorStore

PRECEDENCE_TEXT = (
    "This regional handbook is read together with the global Acme Employee "
    "Handbook. Where a conflict arises specifically with respect to PAID TIME "
    "OFF (PTO), the LOCAL PTO POLICY set out in this APAC Benefits Handbook "
    "TAKES PRECEDENCE over any conflicting PTO provision in the global handbook "
    "for employees covered by this document."
)

PRECEDENCE_JSON = """
{
  "relationships": [
    {
      "target_doc_ref": "global Acme Employee Handbook",
      "relation_type": "precedence",
      "topic": "PTO",
      "precedence": "source_over_target",
      "source_text": "TAKES PRECEDENCE over any conflicting PTO provision"
    }
  ],
  "geographic_scope": null
}
"""

GEO_JSON = (
    '{"relationships": [], "geographic_scope": '
    '{"included": ["Singapore", "Malaysia"], "excluded": []}}'
)


class ScriptedLLMClient:
    """Returns a canned response keyed by a substring of the prompted content --
    lets one test exercise both the relationship and geo-scope extraction paths
    without needing a real LLM."""

    def __init__(self, responses: dict[str, str]) -> None:
        self._responses = responses

    def generate(self, messages: list[Message], stream: bool = True) -> Iterable[str]:
        content = messages[-1].content
        for needle, response in self._responses.items():
            if needle in content:
                yield response
                return
        yield "{}"


def _make_apac_docx(path) -> None:
    document = docx.Document()
    document.add_paragraph("APAC Benefits Handbook", style="Title")
    document.add_paragraph("Conflicts and Precedence", style="Heading 1")
    document.add_paragraph(PRECEDENCE_TEXT)
    document.add_paragraph("Eligibility", style="Heading 2")
    document.add_paragraph("This section applies only to employees in Singapore and Malaysia.")
    document.save(str(path))


async def test_pipeline_applies_relationship_and_geo_override_end_to_end(tmp_path):
    docx_path = tmp_path / "APAC Benefits Handbook - version 2.docx"
    _make_apac_docx(docx_path)

    llm = ScriptedLLMClient({"TAKES PRECEDENCE": PRECEDENCE_JSON, "Singapore and Malaysia": GEO_JSON})
    vector_store = InMemoryVectorStore()
    pipeline = IngestionPipeline(
        chunker=SectionAwareChunker(),
        embedder=FakeEmbedder(),
        vector_store=vector_store,
        metadata_extractor=CompositeMetadataExtractor(llm_client=llm),
        relationship_extractor=SectionAnnotationExtractor(llm_client=llm),
    )

    async with async_session_factory() as db:
        document = Document(
            filename=docx_path.name,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            storage_path=str(docx_path),
        )
        db.add(document)
        await db.commit()
        await db.refresh(document)
        doc_id = document.id

        await pipeline.run(db, doc_id, docx_path)

        await db.refresh(document)
        assert document.status == "ready"
        assert document.title == "APAC Benefits Handbook"  # from filename
        assert document.version == "2"  # from filename
        assert document.is_latest is True

        result = await db.execute(
            select(DocumentRelationship).where(DocumentRelationship.source_doc_id == doc_id)
        )
        relationship = result.scalar_one()
        assert relationship.relation_type == "precedence"
        assert relationship.topic == "PTO"

    # Pull everything back out of the vector store to check per-chunk metadata --
    # the Eligibility section should carry the narrower geo override, not null.
    query_vector = FakeEmbedder().embed(["query"])[0]
    all_chunks = {
        tuple(scored.chunk.metadata["heading_path"]): scored.chunk
        for scored in vector_store.search(query_vector, top_k=10)
    }

    eligibility_chunk = all_chunks[("Conflicts and Precedence", "Eligibility")]
    assert eligibility_chunk.metadata["applicable_regions"] == {
        "included": ["Singapore", "Malaysia"],
        "excluded": [],
    }
    assert eligibility_chunk.metadata["is_latest"] is True

    precedence_chunk = all_chunks[("Conflicts and Precedence",)]
    assert precedence_chunk.metadata["applicable_regions"] is None


PERSONNEL_METADATA_JSON = (
    '{"doc_type": null, "title": null, "effective_date": null, '
    '"applicable_regions": null, '
    '"applicable_personnel": {"included": ["employees"], "excluded": ["contractors"]}}'
)


async def test_pipeline_applies_personnel_scope_to_chunk_metadata(tmp_path):
    docx_path = tmp_path / "Personnel Scope Handbook.docx"
    document = docx.Document()
    document.add_paragraph("Personnel Scope Handbook", style="Title")
    document.add_paragraph(
        "This handbook applies to employees only. It does not apply to contractors, "
        "who should refer to their separate contractor agreement instead."
    )
    document.save(str(docx_path))

    llm = ScriptedLLMClient({"does not apply to contractors": PERSONNEL_METADATA_JSON})
    vector_store = InMemoryVectorStore()
    pipeline = IngestionPipeline(
        chunker=SectionAwareChunker(),
        embedder=FakeEmbedder(),
        vector_store=vector_store,
        metadata_extractor=CompositeMetadataExtractor(llm_client=llm),
        relationship_extractor=SectionAnnotationExtractor(llm_client=llm),
    )
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    async with async_session_factory() as db:
        document_row = Document(filename=docx_path.name, content_type=content_type, storage_path=str(docx_path))
        db.add(document_row)
        await db.commit()
        await db.refresh(document_row)
        await pipeline.run(db, document_row.id, docx_path)
        await db.refresh(document_row)

        assert document_row.status == "ready"
        assert document_row.applicable_personnel == {
            "included": ["employees"],
            "excluded": ["contractors"],
        }
        doc_id = document_row.id

    query_vector = FakeEmbedder().embed(["contractors"])[0]
    chunks = [s for s in vector_store.search(query_vector, top_k=10) if s.chunk.doc_id == doc_id]
    assert chunks
    assert all(c.chunk.metadata["personnel_included"] == ["employees"] for c in chunks)
    assert all(c.chunk.metadata["personnel_excluded"] == ["contractors"] for c in chunks)


def _make_simple_docx(path, text: str) -> None:
    # A short/plain paragraph gets misclassified by `unstructured` as a bare
    # "Title" with nothing under it (sometimes even a short *body* paragraph
    # gets classified as a second Title), which SectionAwareChunker correctly
    # drops as content-free (see test_chunking.py) -- use a real multi-sentence
    # narrative paragraph so it's reliably classified as body text and actually
    # gets chunked/embedded.
    document = docx.Document()
    document.add_paragraph("Regression Versioning Handbook", style="Title")
    document.add_paragraph(text)
    document.save(str(path))


async def test_superseded_document_chunks_get_is_latest_patched(tmp_path):
    # Regression test for a real bug found in end-to-end testing: a document's
    # chunks are embedded with a snapshot of is_latest at THEIR OWN ingestion
    # time. If an older document is ingested before a newer version of the same
    # title exists, its chunks are stamped is_latest=True and -- without the fix
    # in IngestionPipeline._resolve_version -- stay that way forever, even after
    # the newer version arrives and the Document row itself is correctly flipped.
    v1_path = tmp_path / "Regression Versioning Handbook - version 1.docx"
    v2_path = tmp_path / "Regression Versioning Handbook - version 2.docx"
    _make_simple_docx(
        v1_path,
        "This is the first version of the handbook policy. It describes the "
        "standard rules that apply to all employees across the organization.",
    )
    _make_simple_docx(
        v2_path,
        "This is the second version of the handbook policy. It describes the "
        "updated rules that apply to all employees across the organization.",
    )

    vector_store = InMemoryVectorStore()
    llm = FakeLLMClient()
    pipeline = IngestionPipeline(
        chunker=SectionAwareChunker(),
        embedder=FakeEmbedder(),
        vector_store=vector_store,
        metadata_extractor=CompositeMetadataExtractor(llm_client=llm),
        relationship_extractor=SectionAnnotationExtractor(llm_client=llm),
    )
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    async with async_session_factory() as db:
        doc_v1 = Document(filename=v1_path.name, content_type=content_type, storage_path=str(v1_path))
        db.add(doc_v1)
        await db.commit()
        await db.refresh(doc_v1)
        await pipeline.run(db, doc_v1.id, v1_path)
        await db.refresh(doc_v1)
        assert doc_v1.is_latest is True

        doc_v2 = Document(filename=v2_path.name, content_type=content_type, storage_path=str(v2_path))
        db.add(doc_v2)
        await db.commit()
        await db.refresh(doc_v2)
        await pipeline.run(db, doc_v2.id, v2_path)
        await db.refresh(doc_v2)

        await db.refresh(doc_v1)
        assert doc_v1.is_latest is False  # DB row correctly flipped (already worked before the fix)

        v1_doc_id = doc_v1.id

    query_vector = FakeEmbedder().embed(["handbook policy"])[0]
    v1_chunks = [s for s in vector_store.search(query_vector, top_k=10) if s.chunk.doc_id == v1_doc_id]
    assert v1_chunks
    assert all(c.chunk.metadata["is_latest"] is False for c in v1_chunks)
